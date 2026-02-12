"""Export service for generating multi-format output (MD, HTML, PDF, DOCX)."""
import io
import re
import logging
from typing import Optional

import markdown as md_lib
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)

# CSS for standalone HTML and PDF exports
BASE_CSS = """
body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    max-width: 800px; margin: 40px auto; padding: 0 20px;
    color: #1e293b; line-height: 1.7; font-size: 15px;
}
h1 { font-size: 1.8rem; margin: 24px 0 12px; border-bottom: 2px solid #e2e8f0; padding-bottom: 8px; }
h2 { font-size: 1.5rem; margin: 20px 0 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 6px; }
h3 { font-size: 1.25rem; margin: 16px 0 8px; }
h4 { font-size: 1.1rem; margin: 14px 0 6px; }
p { margin: 0 0 12px; }
ul, ol { margin: 0 0 12px; padding-left: 24px; }
li { margin-bottom: 4px; }
blockquote {
    border-left: 4px solid PRIMARY_COLOR;
    margin: 12px 0; padding: 8px 16px;
    background: #f8fafc; color: #64748b;
}
pre {
    background: #1e293b; color: #e2e8f0; padding: 16px; border-radius: 6px;
    overflow-x: auto; margin: 12px 0; font-size: 0.88rem; line-height: 1.5;
}
code {
    background: #f1f5f9; padding: 2px 6px; border-radius: 4px;
    font-size: 0.88rem; font-family: 'Consolas', 'Courier New', monospace;
}
pre code { background: none; padding: 0; color: inherit; }
table { width: 100%; border-collapse: collapse; margin: 12px 0; }
th, td { padding: 8px 12px; border: 1px solid #e2e8f0; text-align: left; }
th { background: #f8fafc; font-weight: 600; }
tr:nth-child(even) { background: #fafbfc; }
a { color: PRIMARY_COLOR; }
hr { border: none; border-top: 2px solid #e2e8f0; margin: 20px 0; }
img { max-width: 100%; }
.header { text-align: center; padding: 16px 0; border-bottom: 2px solid PRIMARY_COLOR; margin-bottom: 24px; color: PRIMARY_COLOR; font-weight: 700; font-size: 1.1rem; }
.footer { text-align: center; padding: 16px 0; border-top: 1px solid #e2e8f0; margin-top: 24px; color: #64748b; font-size: 0.85rem; }
"""


def markdown_to_html(raw: str) -> str:
    """Convert raw markdown to HTML using Python markdown library."""
    extensions = ['tables', 'fenced_code', 'codehilite', 'nl2br', 'toc']
    extension_configs = {
        'codehilite': {'css_class': 'highlight', 'guess_lang': True},
    }
    return md_lib.markdown(raw, extensions=extensions, extension_configs=extension_configs)


def export_markdown(raw: str, title: str = 'Export') -> tuple:
    """Return raw markdown content and filename."""
    filename = _safe_filename(title, 'md')
    return raw.encode('utf-8'), filename, 'text/markdown; charset=utf-8'


def export_html(
    raw: str,
    title: str = 'Export',
    header_text: str = '',
    footer_text: str = '',
    primary_color: str = '#2563eb',
) -> tuple:
    """Generate standalone HTML page with embedded CSS."""
    html_body = markdown_to_html(raw)
    css = BASE_CSS.replace('PRIMARY_COLOR', primary_color)

    header_block = f'<div class="header">{_escape(header_text)}</div>' if header_text else ''
    footer_block = f'<div class="footer">{_escape(footer_text)}</div>' if footer_text else ''

    full_html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{_escape(title)}</title>
    <style>{css}</style>
</head>
<body>
{header_block}
{html_body}
{footer_block}
</body>
</html>"""

    filename = _safe_filename(title, 'html')
    return full_html.encode('utf-8'), filename, 'text/html; charset=utf-8'


def export_pdf(
    raw: str,
    title: str = 'Export',
    header_text: str = '',
    footer_text: str = '',
    primary_color: str = '#2563eb',
) -> tuple:
    """Generate PDF from markdown via xhtml2pdf."""
    from xhtml2pdf import pisa

    html_body = markdown_to_html(raw)
    css = BASE_CSS.replace('PRIMARY_COLOR', primary_color)

    header_block = f'<div class="header">{_escape(header_text)}</div>' if header_text else ''
    footer_block = f'<div class="footer">{_escape(footer_text)}</div>' if footer_text else ''

    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>{css}
    @page {{ size: A4; margin: 2cm; }}
    </style>
</head>
<body>
{header_block}
{html_body}
{footer_block}
</body>
</html>"""

    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(full_html, dest=buffer, encoding='utf-8')

    if pisa_status.err:
        logger.error('PDF generation failed: %s', pisa_status.err)
        raise RuntimeError('Erreur lors de la generation du PDF')

    filename = _safe_filename(title, 'pdf')
    return buffer.getvalue(), filename, 'application/pdf'


def export_docx(
    raw: str,
    title: str = 'Export',
    header_text: str = '',
    footer_text: str = '',
    primary_color: str = '#2563eb',
) -> tuple:
    """Generate Word document from markdown by parsing lines."""
    doc = Document()

    # Parse primary color to RGB
    r, g, b = _hex_to_rgb(primary_color)

    # Header
    if header_text:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(r, g, b)
        doc.add_paragraph()  # spacer

    # Parse markdown line by line
    lines = raw.split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.size = Pt(9)
                run.font.name = 'Courier New'
                p.style = doc.styles['Normal']
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith('#'):
            level = 0
            for ch in stripped:
                if ch == '#':
                    level += 1
                else:
                    break
            text = stripped[level:].strip()
            heading_level = min(level, 4)
            heading = doc.add_heading(text, level=heading_level)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(r, g, b)
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(100, 116, 139)
            continue

        # Unordered list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            continue

        # Ordered list
        ol_match = re.match(r'^\d+\.\s+(.*)', stripped)
        if ol_match:
            doc.add_paragraph(ol_match.group(1), style='List Number')
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.add_run('_' * 60)
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_inline_formatting(p, stripped)

    # Footer
    if footer_text:
        doc.add_paragraph()  # spacer
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)

    buffer = io.BytesIO()
    doc.save(buffer)
    filename = _safe_filename(title, 'docx')
    return buffer.getvalue(), filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def _add_inline_formatting(paragraph, text: str):
    """Parse inline markdown (bold, italic, code) and add runs to paragraph."""
    # Simple inline parsing: **bold**, *italic*, `code`
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))'
    for match in re.finditer(pattern, text):
        full = match.group(0)
        if full.startswith('**') and full.endswith('**'):
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif full.startswith('*') and full.endswith('*'):
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif full.startswith('`') and full.endswith('`'):
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(match.group(5))


def _safe_filename(title: str, ext: str) -> str:
    """Create safe filename from title."""
    safe = re.sub(r'[^\w\s-]', '', title).strip()
    safe = re.sub(r'[\s]+', '_', safe)
    if not safe:
        safe = 'export'
    return f'{safe[:80]}.{ext}'


def _escape(text: str) -> str:
    """Escape HTML characters."""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;'))


def _hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    if len(hex_color) != 6:
        return (37, 99, 235)  # default blue
    try:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    except ValueError:
        return (37, 99, 235)
